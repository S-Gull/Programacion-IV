from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def crear_documento_dom():
    # Crear un nuevo documento
    doc = Document()

    # Configuración de estilo
    estilo = doc.styles['Normal']
    fuente = estilo.font
    fuente.name = 'Calibri'
    fuente.size = Pt(11)

    # Título principal
    doc.add_heading('El DOM (Document Object Model) en JavaScript', level=0).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Sección 1: ¿Qué es el DOM?
    doc.add_heading('¿Qué es el DOM?', level=1)
    doc.add_paragraph(
        "El DOM (Document Object Model) es una interfaz de programación para documentos HTML y XML. "
        "Representa la estructura del documento como un árbol de nodos, donde cada nodo es un objeto que "
        "representa una parte del documento (elementos, atributos, texto, etc.)."
    )
    doc.add_paragraph(
        "El DOM permite a los desarrolladores manipular el contenido, la estructura y el estilo de una "
        "página web dinámicamente usando JavaScript."
    )

    # Sección 2: Características del DOM
    doc.add_heading('Características del DOM', level=1)
    caracteristicas = [
        "Estructura de Árbol: El DOM representa el documento como un árbol de nodos.",
        "Manipulación Dinámica: Permite agregar, eliminar o modificar elementos y atributos.",
        "Acceso a Elementos: Proporciona métodos como getElementById, querySelector, etc.",
        "Eventos: Permite manejar interacciones del usuario (clics, teclas, etc.).",
        "Independencia de Plataforma: Funciona en cualquier navegador moderno."
    ]
    for item in caracteristicas:
        doc.add_paragraph(item, style='List Bullet')

    # Sección 3: Métodos Comunes del DOM
    doc.add_heading('Métodos Comunes del DOM', level=1)
    metodos = [
        ("Seleccionar Elementos", "document.getElementById('id'), document.querySelector('.clase')"),
        ("Crear y Agregar Elementos", "document.createElement('div'), elemento.appendChild(nuevoElemento)"),
        ("Modificar Contenido y Estilos", "elemento.textContent = 'Texto', elemento.style.color = 'red'"),
        ("Manejo de Eventos", "elemento.addEventListener('click', función)")
    ]
    for titulo, descripcion in metodos:
        doc.add_heading(titulo, level=2)
        doc.add_paragraph(descripcion)

    # Sección 4: Propiedades Comunes del DOM
    doc.add_heading('Propiedades Comunes del DOM', level=1)
    propiedades = [
        "innerHTML: Obtiene o establece el contenido HTML de un elemento.",
        "textContent: Obtiene o establece el texto de un elemento.",
        "parentNode: Devuelve el nodo padre de un elemento.",
        "childNodes: Devuelve una lista de los nodos hijos.",
        "classList: Permite manipular las clases CSS de un elemento."
    ]
    for item in propiedades:
        doc.add_paragraph(item, style='List Bullet')

    # Sección 5: Ejemplos Prácticos
    doc.add_heading('Ejemplos Prácticos', level=1)
    ejemplos = [
        ("Seleccionar Elementos", "document.getElementById('miElemento').textContent = '¡Hola!';"),
        ("Crear y Agregar Elementos", "const nuevoElemento = document.createElement('p'); nuevoElemento.textContent = 'Nuevo párrafo'; document.body.appendChild(nuevoElemento);"),
        ("Modificar Estilos", "document.getElementById('miElemento').style.color = 'blue';"),
        ("Manejo de Eventos", "document.getElementById('miBoton').addEventListener('click', () => { alert('¡Hola!'); });")
    ]
    for titulo, codigo in ejemplos:
        doc.add_heading(titulo, level=2)
        doc.add_paragraph(codigo, style='Intense Quote')

    # Sección 6: Conclusión
    doc.add_heading('Conclusión', level=1)
    doc.add_paragraph(
        "El DOM es una herramienta fundamental para la manipulación dinámica de páginas web. "
        "Permite a los desarrolladores interactuar con los elementos HTML, modificar su contenido, "
        "estilo y comportamiento en respuesta a las acciones del usuario. Dominar el DOM es esencial "
        "para crear aplicaciones web interactivas y modernas."
    )

    # Guardar el documento
    doc.save('./DOM_JavaScript.docx')
    print("Documento 'DOM_JavaScript.docx' generado con éxito.")

# Ejecutar la función para crear el documento
crear_documento_dom()
